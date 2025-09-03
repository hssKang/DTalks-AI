import os
import uuid
import base64
import logging
import requests
import pdfplumber
import pandas as pd
import tempfile
from src.layers.LLM.bedrock_model import call_image_ocr

from io import BytesIO
from docx import Document
from dotenv import load_dotenv
from qdrant_client.models import PointStruct
from src.utils.database.connect_qdrant import init_qdrant

from src.utils.tools.embedding import vectorize


# 환경 변수 설정
env_path = os.path.join(os.path.dirname(__file__), "..", "..", "..", ".env")
load_dotenv(dotenv_path=os.path.abspath(env_path))

QDRANT_HOST = os.getenv("QDRANT_HOST")
QDRANT_PORT = int(os.getenv("QDRANT_PORT"))
QDRANT_COLLECTION = "internal_documents"
qdrant_client = None

#############################################################################


# Qdrant에 데이터 저장
def save_data(file_name, ori, vector):
    global qdrant_client
    # Qdrant 클라이언트 lazy initialization (이중 체크)
    if qdrant_client is None:
        qdrant_client = init_qdrant(QDRANT_COLLECTION)
    point_id = str(uuid.uuid4())
    points = [
        PointStruct(
            id=point_id, vector=vector, payload={"text": ori, "file_name": file_name}
        )
    ]
    qdrant_client.upsert(collection_name=QDRANT_COLLECTION, points=points)


##########################################################################


# 문단 단위로 chbunking
def chunk_by_paragraph(file, ext=None, max_tokens=200):
    # 파일 확장자 확인
    if isinstance(file, str):
        document = os.path.splitext(file)[-1].lower()
    elif isinstance(file, BytesIO):
        if ext is None:
            logging.error("BytesIO 객체는 확장자를 별도로 전달해야 합니다.")
            raise ValueError("BytesIO 객체는 확장자를 별도로 전달해야 합니다.")
        document = ext
    else:
        logging.error("file은 str 또는 BytesIO 객체여야 합니다.")
        raise ValueError("file은 str 또는 BytesIO 객체여야 합니다.")

    chunks = []
    current_chunk = ""

    # docx 파일에서 텍스트 추출
    if document == ".docx":
        doc = Document(file)
        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                logging.warning("빈 문단이 발견되었습니다. 무시합니다.")
                continue

            is_heading = para.style.name.startswith("Heading")

            if is_heading:
                if current_chunk:
                    chunks.append(current_chunk.strip())
                current_chunk = text + "\n"
            else:
                # 현재 청크에 문단을 추가할 수 있는지 확인
                if len(current_chunk) + len(text) < max_tokens:
                    current_chunk += text + "\n"
                else:
                    # 길이 초과 시 현재 청크를 저장하고 새 청크 시작
                    chunks.append(current_chunk.strip())
                    current_chunk = text + "\n"

    # pdf 파일에서 텍스트 추출
    elif document == ".pdf":
        full_text = ""
        # 모든 페이지를 순회하며 텍스트 추출
        with pdfplumber.open(file) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    full_text += page_text + "\n"
        paragraphs = [p.strip() for p in full_text.split("\n\n") if p.strip()]
        for para in paragraphs:
            if len(current_chunk) + len(para) < max_tokens:
                current_chunk += para + "\n"
            else:
                chunks.append(current_chunk.strip())
                current_chunk = para + "\n"

    else:
        logging.error(f"지원하지 않는 파일 형식입니다: {document}")

    # 마지막 청크가 남아 있다면 추가
    if current_chunk:
        chunks.append(current_chunk.strip())

    return chunks


# 슬라이딩 윈도우 방식으로 텍스트를 청크 단위로 나누기
def chunk_by_sliding_window(text, chunk_size=200, overlap=50):
    chunks = []
    start = 0
    while start < len(text):
        end = start + chunk_size
        chunks.append(text[start:end])
        start += chunk_size - overlap
    return chunks


# 엑셀 파일에서 텍스트 추출
def chunk_xlsx(file_path, sheet_name=0):
    df = pd.read_excel(file_path, sheet_name=sheet_name)
    # 문자열 컬럼만 ""로 채움
    df[df.select_dtypes(include=["object"]).columns] = df.select_dtypes(
        include=["object"]
    ).fillna("")

    # 숫자형 컬럼은 0으로 채움 (또는 NaN 그대로 유지)
    df[df.select_dtypes(include=["number"]).columns] = df.select_dtypes(
        include=["number"]
    ).fillna(0)

    chunks = []
    # 각 행을 하나의 청크로 저장
    for i in range(len(df)):
        row = df.iloc[i]
        row_text = " | ".join([str(cell) for cell in row])
        chunks.append(row_text.strip())
    return chunks


# CSV 파일에서 텍스트 추출
def chunk_csv(file_path):
    df = pd.read_csv(file_path)
    df.fillna("", inplace=True)
    chunks = []
    # 각 행을 하나의 청크로 저장
    for i in range(len(df)):
        row = df.iloc[i]
        row_text = " | ".join([str(cell) for cell in row])
        chunks.append(row_text.strip())
    return chunks


########################################################################


# 이미지 파일에서 OCR을 통해 텍스트 추출 후 청크
def chunk_image_ocr(file_path, max_tokens=200):
    try:
        # 이미지 파일을 base64로 인코딩
        with open(file_path, "rb") as image_file:
            image_data = image_file.read()
            image_base64 = base64.b64encode(image_data).decode("utf-8")

        # 이미지 확장자 확인
        file_extension = os.path.splitext(file_path)[-1].lower()
        if file_extension in [".jpg", ".jpeg"]:
            media_type = "image/jpeg"
        elif file_extension == ".png":
            media_type = "image/png"
        else:
            logging.error(f"지원하지 않는 이미지 형식입니다: {file_extension}")
            raise ValueError(f"지원하지 않는 이미지 형식입니다: {file_extension}")

        # bedrock_model.py의 call_image_ocr 함수 사용
        extracted_text = call_image_ocr(image_base64, media_type)

        # 슬라이딩 윈도우 방식으로 청크 분할
        chunks = chunk_by_sliding_window(
            extracted_text, chunk_size=max_tokens, overlap=50
        )
        return chunks

    except Exception as e:
        logging.error(f"이미지 OCR 처리 중 오류 발생: {str(e)}")
        raise RuntimeError(f"이미지 OCR 처리 중 오류 발생: {str(e)}")


########################################################################


# 파일 처리 및 저장
def process_and_store(file_info):
    global qdrant_client
    # Qdrant 클라이언트 lazy initialization
    if qdrant_client is None:
        qdrant_client = init_qdrant(QDRANT_COLLECTION)

    file_url = file_info["fileUrl"]
    description = file_info.get("description", "")
    file_name = os.path.basename(file_url)
    document = os.path.splitext(file_name)[-1].lower()

    # 파일 다운로드
    response = requests.get(file_url)
    if response.status_code != 200:
        logging.error(f"파일 다운로드 실패: {file_url}")
        return

    chunks = []
    # 이미지 파일 처리
    if document in [".png", ".jpg", ".jpeg"]:
        try:
            # 임시 파일로 저장
            with tempfile.NamedTemporaryFile(
                delete=False, suffix=document
            ) as temp_file:
                temp_file.write(response.content)
                temp_file_path = temp_file.name

            # AWS Bedrock을 사용한 OCR
            ocr_chunks = chunk_image_ocr(temp_file_path)

            # 설명과 OCR 결과 합치기
            if description.strip():
                final_text = f"{description.strip()}\n\n" + "\n".join(ocr_chunks)
                chunks = chunk_by_sliding_window(final_text, chunk_size=200, overlap=50)
            else:
                chunks = ocr_chunks

            # 임시 파일 삭제
            os.unlink(temp_file_path)

        except Exception as e:
            logging.error(f"이미지 처리 중 오류 발생: {str(e)}")
            return

    # DOCX 파일 처리
    elif document == ".docx":
        doc = Document(BytesIO(response.content))
        doc_text = "\n".join(
            [para.text.strip() for para in doc.paragraphs if para.text.strip()]
        )
        final_text = (
            f"{description}\n{doc_text}".strip() if description else doc_text.strip()
        )
        if final_text:
            chunks = chunk_by_paragraph(BytesIO(response.content), ext=document)
    # PDF 파일 처리
    elif document == ".pdf":
        chunks = chunk_by_paragraph(BytesIO(response.content), ext=document)
        if len(chunks) <= 1 or any(len(c) > 200 for c in chunks):
            with pdfplumber.open(BytesIO(response.content)) as pdf:
                text = ""
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
            final_text = (
                f"{description}\n{text}".strip() if description else text.strip()
            )
            if not final_text:
                logging.error("[PDF] 텍스트 추출 실패, OCR 시도 필요")
                logging.warning("PDF에서 텍스트 추출 실패, pdf가 이미지 형태로 저장됨")
            # 슬라이딩 윈도우 방식으로 청크
            if final_text:
                logging.info("[PDF] 슬라이딩 윈도우 방식으로 재청크")
                chunks = chunk_by_sliding_window(final_text, chunk_size=200, overlap=50)
                if chunks == 0:
                    logging.warning(
                        "슬라이딩 윈도우 방식으로 청크가 생성되지 않았습니다."
                    )
    # XLSX 파일 처리
    elif document == ".xlsx":
        chunks = chunk_xlsx(BytesIO(response.content))

    # CSV 파일 처리
    elif document == ".csv":
        chunks = chunk_csv(BytesIO(response.content))
    else:
        logging.error(f"지원하지 않는 파일 형식입니다: {document}")
        return

    logging.info(f"[{document.upper()}] 청크 개수: {len(chunks)}")
    for i, chunk in enumerate(chunks):
        if not chunk.strip():
            logging.warning(f"빈 청크 발견, 벡터화 및 저장을 건너뜁니다. (index: {i})")
            continue
        logging.info(f"--- 청크 {i+1} ---\n{chunk}\n")
        # 임베딩 및 Qdrant 저장
        vector = vectorize(chunk)
        save_data(file_name, chunk, vector)


if __name__ == "__main__":
    qdrant_client = init_qdrant(QDRANT_COLLECTION)

    file_info_list = [
        {
            "fileUrl": "https://nwmmaxrrqgkmeshqdjvu.supabase.co/storage/v1/object/public/roadmapdb//file_example.docx",
        },
        {
            "fileUrl": "https://nwmmaxrrqgkmeshqdjvu.supabase.co/storage/v1/object/public/roadmapdb//file_example.xlsx",
        },
        {
            "fileUrl": "https://nwmmaxrrqgkmeshqdjvu.supabase.co/storage/v1/object/public/roadmapdb//file_sample.pdf",
        },
        {
            "fileUrl": "https://nwmmaxrrqgkmeshqdjvu.supabase.co/storage/v1/object/public/roadmapdb//sa.png",
            "description": """
    1. Local에서 frontend와 backend가 각각 develop 브랜치에 push를 하면, Github Actions가 자동 감지하여 Event trigger를 발생시킨다.
    2. 이때,  Github Actions에서 프론트엔드, 백엔드 파일을 자체 빌드하고, 압축파일 형태로 만든다.
    3. 그리고 S3 bucket으로 zip 파일을 전송시키고, AWS IAM에서 권한 설정에 따라 CodeDeploy를 통해 zip 파일을 Amazon Ec2에 압축을 풀게 된다.
    4. EC2는 가비아에서 구매한 moamoadev.shop 도메인을 설정하였는데, AWS에서 지원하는 Route53을 통해 도메인을 등록 시켜주었다.
    5. 무중단 배포와 리버스 프록시 설정을 위해 Amazon EC2에 Nginx를 설치하고, Nginx 내에서 리버스 프록시를 설정하여, / 요청은 3000번 포트(리액트 정적 파일), /api 요청은 8080 포트(스프링부트 내장 톰캣)으로 리다이렉트하게 만들었다.
    6. SSL Certificate(Https) 설정은 Let's encrypt를 사용하여 무료로 자동으로 업데이트되도록 https 설정을 진행하였다. 
    7. 백엔드내의 기술스택은 QueryDSL, redis(refreshToken 저장을 위한 내장 저장소), Spring Data JPA, Spring Security를 사용하였다. 백엔드에서는 DB는 AmazonRDS를 사용하고 있으며, MySQL을 이용하고 있다.
    8. Spring Security를 통해 OAuth 2.0(카카오, 네이버, 구글)을 자동으로 백엔드에서 처리후 프론트로 리다이렉트하게 설정하였다. JWT를 이용하여, accessToken, refreshToken(보안을 위해 쿠키로 전달)로 진행하였음.
    """,
        },
    ]
    # 받아온 파일 정보로 Qdrant에 저장
    for file_info in file_info_list:
        logging.info(f"Qdrant에 저장 중: {os.path.basename(file_info['fileUrl'])}")
        process_and_store(file_info)
